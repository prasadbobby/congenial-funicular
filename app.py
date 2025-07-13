from flask import Flask, render_template, request, jsonify, send_from_directory, session, flash, redirect, url_for
from flask_cors import CORS
import pandas as pd
import os
import numpy as np
import torch
import io
from openai import AzureOpenAI
import json
import time
import requests
import markdown2
from PIL import Image
import threading
import logging
import queue
import re
import uuid
import random
import string
import base64
import httpx
from azure.communication.email import EmailClient
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple, Union
from dotenv import load_dotenv
import PyPDF2
import docx
import tempfile
from werkzeug.utils import secure_filename
import matplotlib.pyplot as plt
import seaborn as sns
import hashlib
import fitz  # PyMuPDF for better PDF handling
from docx import Document
import csv
import xlrd
import openpyxl
from pathlib import Path
import plotly.graph_objects as go
import plotly.express as px
import plotly.utils
from io import BytesIO
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

try:
    import google.generativeai as genai
    GOOGLE_AI_AVAILABLE = True
except ImportError:
    GOOGLE_AI_AVAILABLE = False
    genai = None

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("app.log")
    ]
)
logger = logging.getLogger("mediassist_agent")

# Initialize Flask app
app = Flask(__name__, static_folder='static')
CORS(app)
app.secret_key = os.getenv('SECRET_KEY', 'appointment_booking_secret_key')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Configure upload settings
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'reports'
CHARTS_FOLDER = 'static/charts'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'xlsx', 'xls', 'csv'}

for folder in [UPLOAD_FOLDER, REPORTS_FOLDER, CHARTS_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORTS_FOLDER'] = REPORTS_FOLDER
app.config['CHARTS_FOLDER'] = CHARTS_FOLDER

# Environment variables
AZURE_OPENAI_API_KEY = os.getenv('AZURE_OPENAI_API_KEY')
AZURE_OPENAI_ENDPOINT = os.getenv('AZURE_OPENAI_ENDPOINT')
AZURE_OPENAI_API_VERSION = os.getenv('AZURE_OPENAI_API_VERSION', '2024-02-01')
AZURE_OPENAI_DEPLOYMENT_NAME = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME', 'gpt-4o-mini')
AZURE_OPENAI_MODEL_NAME = os.getenv('AZURE_OPENAI_MODEL_NAME', 'gpt-4o-mini')
AZURE_OPENAI_EMBEDDING_DEPLOYMENT = os.getenv('AZURE_OPENAI_EMBEDDING_DEPLOYMENT', 'text-embedding-3-small')
AZURE_OPENAI_EMBEDDING_MODEL = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL', 'text-embedding-3-small')
AZURE_COMMUNICATION_CONNECTION_STRING = os.getenv('AZURE_COMMUNICATION_CONNECTION_STRING')
AZURE_COMMUNICATION_SENDER_EMAIL = os.getenv('AZURE_COMMUNICATION_SENDER_EMAIL')

# Google AI for image analysis fallback (internal use only)
GOOGLE_AI_API_KEY = os.getenv('GOOGLE_AI_API_KEY')

# Initialize Azure OpenAI client
azure_openai_client = None
if AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT:
    try:
        import httpx
        # Create HTTP client with SSL verification disabled for corporate environments
        http_client = httpx.Client(verify=False)
        
        azure_openai_client = AzureOpenAI(
            api_key=AZURE_OPENAI_API_KEY,
            api_version=AZURE_OPENAI_API_VERSION,
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            http_client=http_client
        )
        logger.info("Azure OpenAI client initialized successfully (SSL verification disabled)")
    except Exception as e:
        logger.error(f"Failed to initialize Azure OpenAI client: {e}")
        azure_openai_client = None

# Initialize Google AI for image analysis fallback (internal use only)
google_ai_model = None
if GOOGLE_AI_API_KEY and GOOGLE_AI_AVAILABLE:
    try:
        genai.configure(api_key=GOOGLE_AI_API_KEY)
        google_ai_model = genai.GenerativeModel('gemini-1.5-flash')
        logger.info("Google AI model initialized successfully for image analysis fallback")
    except Exception as e:
        logger.warning(f"Failed to initialize Google AI model: {e}")
        google_ai_model = None

# Store appointments
appointments = []

# Normal health value ranges
NORMAL_RANGES = {
    'blood_pressure_systolic': {'min': 90, 'max': 120, 'unit': 'mmHg', 'name': 'Systolic Blood Pressure'},
    'blood_pressure_diastolic': {'min': 60, 'max': 80, 'unit': 'mmHg', 'name': 'Diastolic Blood Pressure'},
    'heart_rate': {'min': 60, 'max': 100, 'unit': 'bpm', 'name': 'Heart Rate'},
    'temperature': {'min': 97.0, 'max': 99.5, 'unit': '°F', 'name': 'Body Temperature'},
    'glucose_fasting': {'min': 70, 'max': 100, 'unit': 'mg/dL', 'name': 'Fasting Glucose'},
    'glucose_random': {'min': 70, 'max': 140, 'unit': 'mg/dL', 'name': 'Random Glucose'},
    'cholesterol_total': {'min': 125, 'max': 200, 'unit': 'mg/dL', 'name': 'Total Cholesterol'},
    'cholesterol_ldl': {'min': 0, 'max': 100, 'unit': 'mg/dL', 'name': 'LDL Cholesterol'},
    'cholesterol_hdl_male': {'min': 40, 'max': 999, 'unit': 'mg/dL', 'name': 'HDL Cholesterol (Male)'},
    'cholesterol_hdl_female': {'min': 50, 'max': 999, 'unit': 'mg/dL', 'name': 'HDL Cholesterol (Female)'},
    'triglycerides': {'min': 0, 'max': 150, 'unit': 'mg/dL', 'name': 'Triglycerides'},
    'hemoglobin_male': {'min': 13.8, 'max': 17.2, 'unit': 'g/dL', 'name': 'Hemoglobin (Male)'},
    'hemoglobin_female': {'min': 12.1, 'max': 15.1, 'unit': 'g/dL', 'name': 'Hemoglobin (Female)'},
    'hematocrit_male': {'min': 40.7, 'max': 50.3, 'unit': '%', 'name': 'Hematocrit (Male)'},
    'hematocrit_female': {'min': 36.1, 'max': 44.3, 'unit': '%', 'name': 'Hematocrit (Female)'},
    'white_blood_cells': {'min': 3.5, 'max': 10.5, 'unit': '×10³/μL', 'name': 'White Blood Cells'},
    'platelets': {'min': 150, 'max': 450, 'unit': '×10³/μL', 'name': 'Platelets'},
    'creatinine_male': {'min': 0.74, 'max': 1.35, 'unit': 'mg/dL', 'name': 'Creatinine (Male)'},
    'creatinine_female': {'min': 0.59, 'max': 1.04, 'unit': 'mg/dL', 'name': 'Creatinine (Female)'},
    'bun': {'min': 6, 'max': 24, 'unit': 'mg/dL', 'name': 'Blood Urea Nitrogen'},
    'sodium': {'min': 136, 'max': 145, 'unit': 'mEq/L', 'name': 'Sodium'},
    'potassium': {'min': 3.5, 'max': 5.2, 'unit': 'mEq/L', 'name': 'Potassium'},
    'calcium': {'min': 8.5, 'max': 10.2, 'unit': 'mg/dL', 'name': 'Calcium'},
    'albumin': {'min': 3.5, 'max': 5.0, 'unit': 'g/dL', 'name': 'Albumin'},
    'bilirubin_total': {'min': 0.1, 'max': 1.2, 'unit': 'mg/dL', 'name': 'Total Bilirubin'},
    'alt': {'min': 7, 'max': 56, 'unit': 'U/L', 'name': 'ALT (Alanine Transaminase)'},
    'ast': {'min': 10, 'max': 40, 'unit': 'U/L', 'name': 'AST (Aspartate Transaminase)'},
    'alkaline_phosphatase': {'min': 44, 'max': 147, 'unit': 'U/L', 'name': 'Alkaline Phosphatase'},
    'vitamin_d': {'min': 30, 'max': 100, 'unit': 'ng/mL', 'name': 'Vitamin D'},
    'vitamin_b12': {'min': 200, 'max': 900, 'unit': 'pg/mL', 'name': 'Vitamin B12'},
    'iron': {'min': 60, 'max': 170, 'unit': 'μg/dL', 'name': 'Iron'},
    'ferritin_male': {'min': 12, 'max': 300, 'unit': 'ng/mL', 'name': 'Ferritin (Male)'},
    'ferritin_female': {'min': 12, 'max': 150, 'unit': 'ng/mL', 'name': 'Ferritin (Female)'},
    'tsh': {'min': 0.27, 'max': 4.2, 'unit': 'μIU/mL', 'name': 'TSH (Thyroid Stimulating Hormone)'},
    'hba1c': {'min': 4.0, 'max': 5.6, 'unit': '%', 'name': 'HbA1c (Hemoglobin A1c)'},
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_teams_meeting_id():
    # Microsoft Teams uses a longer format with numbers and letters
    # Example: 19:meeting_NTM3ZWU4ZDUtZTU5ZC00MGY2LTljOGEtOGQ4OWY4Y2Y4NWE2@thread.v2
    meeting_id = ''.join(random.choices(string.ascii_letters + string.digits, k=64))
    return f"19:meeting_{meeting_id}@thread.v2"

def send_appointment_email(appointment, doctor_name, patient_email, meet_link):
    """Send appointment confirmation email using Azure Communication Services"""
    try:
        if not AZURE_COMMUNICATION_CONNECTION_STRING:
            logger.error("Azure Communication Services connection string not configured")
            return False
        
        # Create the email client
        email_client = EmailClient.from_connection_string(AZURE_COMMUNICATION_CONNECTION_STRING)
        
        # Email body
        start_time = appointment['start_time']
        end_time = appointment['end_time']
        
        # Format dates if they're strings
        if isinstance(start_time, str):
            start_time = datetime.fromisoformat(start_time)
        if isinstance(end_time, str):
            end_time = datetime.fromisoformat(end_time)
        
        html_body = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; border: 1px solid #ddd; border-radius: 10px; }}
                .header {{ color: #0078d4; text-align: center; margin-bottom: 20px; }}
                .info-box {{ background-color: #f9f9f9; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                .info-item {{ margin-bottom: 10px; }}
                .info-label {{ font-weight: bold; }}
                .footer {{ margin-top: 20px; color: #666; }}
            </style>
        </head>
        <body>
            <div class="container">
                <h2 class="header">🩺 Your Medical Appointment has been Confirmed</h2>
                <div class="info-box">
                    <div class="info-item">
                        <span class="info-label">Doctor:</span> {doctor_name}
                    </div>
                    <div class="info-item">
                        <span class="info-label">Date:</span> {start_time.strftime("%A, %B %d, %Y")}
                    </div>
                    <div class="info-item">
                        <span class="info-label">Time:</span> {start_time.strftime("%I:%M %p")} - {end_time.strftime("%I:%M %p")}
                    </div>
                    <div class="info-item">
                        <span class="info-label">Teams Meeting Link:</span> <a href="{meet_link}" style="color: #0078d4;">{meet_link}</a>
                    </div>
                </div>
                <p>Please join the meeting at the scheduled time using the meeting link above.</p>
                <div class="footer">
                    <p>Best regards,<br>
                    MediAssist AI - Medical Support System</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Plain text version
        text_body = f"""
        Your Medical Appointment has been Confirmed
        
        Doctor: {doctor_name}
        Date: {start_time.strftime("%A, %B %d, %Y")}
        Time: {start_time.strftime("%I:%M %p")} - {end_time.strftime("%I:%M %p")}
        Teams Meeting Link: {meet_link}
        
        Please join the meeting at the scheduled time using the meeting link above.
        
        Best regards,
        MediAssist AI - Medical Support System
        """
        
        # Create email message
        message = {
            "senderAddress": AZURE_COMMUNICATION_SENDER_EMAIL,
            "recipients": {
                "to": [{"address": patient_email}]
            },
            "content": {
                "subject": f"Medical Appointment Confirmation with {doctor_name}",
                "plainText": text_body,
                "html": html_body
            }
        }
        
        # Send the email
        poller = email_client.begin_send(message)
        result = poller.result()
        
        # Handle result properly (result is a dict, not an object)
        if isinstance(result, dict):
            message_id = result.get('id', 'Unknown')
        else:
            message_id = getattr(result, 'id', 'Unknown')
            
        logger.info(f"Email sent successfully to {patient_email}. Message ID: {message_id}")
        return True
        
    except Exception as e:
        logger.error(f"Error sending email: {str(e)}")
        return False

# Document parsing utilities
def extract_text_from_pdf(file_path):
    """Extract text from PDF file using PyMuPDF"""
    try:
        text = ""
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        return ""

def extract_data_from_excel(file_path):
    """Extract data from Excel file"""
    try:
        # Try reading with openpyxl first
        df = pd.read_excel(file_path, engine='openpyxl')
        return df.to_string()
    except Exception as e:
        try:
            # Fallback to xlrd for older Excel files
            df = pd.read_excel(file_path, engine='xlrd')
            return df.to_string()
        except Exception as e2:
            logger.error(f"Error extracting data from Excel: {str(e2)}")
            return ""

def extract_data_from_csv(file_path):
    """Extract data from CSV file"""
    try:
        df = pd.read_csv(file_path)
        return df.to_string()
    except Exception as e:
        logger.error(f"Error extracting data from CSV: {str(e)}")
        return ""

def extract_health_values(text):
    """Extract health values and measurements from text"""
    health_values = {}
    
    # Common patterns for health values
    patterns = {
        'blood_pressure': r'(?:blood\s*pressure|bp)[\s:]*(\d{2,3})/(\d{2,3})',
        'heart_rate': r'(?:heart\s*rate|pulse|hr)[\s:]*(\d{2,3})',
        'temperature': r'(?:temperature|temp)[\s:]*(\d{2,3}\.?\d*)',
        'glucose': r'(?:glucose|sugar|blood\s*sugar)[\s:]*(\d{2,3})',
        'cholesterol': r'(?:cholesterol|chol)[\s:]*(\d{2,3})',
        'hemoglobin': r'(?:hemoglobin|hb|hgb)[\s:]*(\d{1,2}\.?\d*)',
        'hematocrit': r'(?:hematocrit|hct)[\s:]*(\d{1,2}\.?\d*)',
        'creatinine': r'(?:creatinine|creat)[\s:]*(\d{1}\.?\d*)',
        'bun': r'(?:bun|urea)[\s:]*(\d{1,2})',
        'sodium': r'(?:sodium|na)[\s:]*(\d{2,3})',
        'potassium': r'(?:potassium|k)[\s:]*(\d{1}\.?\d*)',
        'hba1c': r'(?:hba1c|a1c|hemoglobin\s*a1c)[\s:]*(\d{1}\.?\d*)',
        'tsh': r'(?:tsh|thyroid)[\s:]*(\d{1}\.?\d*)',
    }
    
    text_lower = text.lower()
    
    for key, pattern in patterns.items():
        matches = re.finditer(pattern, text_lower, re.IGNORECASE)
        for match in matches:
            if key == 'blood_pressure':
                systolic = int(match.group(1))
                diastolic = int(match.group(2))
                health_values['blood_pressure_systolic'] = systolic
                health_values['blood_pressure_diastolic'] = diastolic
            else:
                try:
                    value = float(match.group(1))
                    health_values[key] = value
                except ValueError:
                    continue
    
    return health_values

def generate_health_chart(health_values, gender='unknown'):
    """Generate visual charts for health values"""
    if not health_values:
        return None
    
    # Create a figure with subplots
    fig = go.Figure()
    
    abnormal_values = []
    chart_data = []
    
    for key, value in health_values.items():
        # Find appropriate normal range
        range_key = key
        if key in ['cholesterol_hdl', 'hemoglobin', 'hematocrit', 'creatinine', 'ferritin']:
            if gender.lower() == 'male':
                range_key = f"{key}_male"
            elif gender.lower() == 'female':
                range_key = f"{key}_female"
        
        if range_key in NORMAL_RANGES:
            normal_range = NORMAL_RANGES[range_key]
            min_val = normal_range['min']
            max_val = normal_range['max']
            name = normal_range['name']
            unit = normal_range['unit']
            
            # Determine status
            if value < min_val:
                status = 'Low'
                color = '#ff4444'
            elif value > max_val:
                status = 'High'
                color = '#ff4444'
            else:
                status = 'Normal'
                color = '#44ff44'
            
            if status != 'Normal':
                abnormal_values.append({
                    'name': name,
                    'value': value,
                    'unit': unit,
                    'status': status,
                    'normal_range': f"{min_val}-{max_val} {unit}"
                })
            
            chart_data.append({
                'name': name,
                'value': value,
                'min': min_val,
                'max': max_val,
                'status': status,
                'color': color,
                'unit': unit
            })
    
    if not chart_data:
        return None
    
    # Create bar chart
    names = [item['name'] for item in chart_data]
    values = [item['value'] for item in chart_data]
    colors = [item['color'] for item in chart_data]
    
    fig.add_trace(go.Bar(
        x=names,
        y=values,
        marker_color=colors,
        text=[f"{item['value']} {item['unit']}" for item in chart_data],
        textposition='outside',
        name='Your Values'
    ))
    
    # Add normal range indicators
    for i, item in enumerate(chart_data):
        fig.add_shape(
            type="line",
            x0=i-0.4,
            y0=item['min'],
            x1=i+0.4,
            y1=item['min'],
            line=dict(color="blue", width=2, dash="dash"),
        )
        fig.add_shape(
            type="line",
            x0=i-0.4,
            y0=item['max'],
            x1=i+0.4,
            y1=item['max'],
            line=dict(color="blue", width=2, dash="dash"),
        )
    
    fig.update_layout(
        title="Health Values Analysis",
        xaxis_title="Health Parameters",
        yaxis_title="Values",
        xaxis_tickangle=-45,
        height=600,
        showlegend=False,
        annotations=[
            dict(
                text="Blue dashed lines indicate normal range",
                xref="paper", yref="paper",
                x=0.5, y=1.02, xanchor='center', yanchor='bottom',
                showarrow=False,
                font=dict(size=12, color="blue")
            )
        ]
    )
    
    # Save chart
    chart_filename = f"health_chart_{uuid.uuid4().hex[:8]}.html"
    chart_path = os.path.join(app.config['CHARTS_FOLDER'], chart_filename)
    fig.write_html(chart_path)
    
    return {
        'chart_path': f"/static/charts/{chart_filename}",
        'abnormal_values': abnormal_values,
        'total_values': len(chart_data)
    }

# Define the agent system architecture
class MediAssistAgentSystem:
    def __init__(self):
        self.agents = {}
        self.model_manager = ModelManager()
        self.data_manager = DataManager(self.model_manager)
        self.task_queue = queue.Queue()
        self.results_cache = {}
        self.is_initialized = False
        self.embedding_cache = {}

    def initialize(self):
        """Initialize the agent system with all required components"""
        logger.info("Initializing MediAssist Agent System")
        
        # Initialize model manager
        self.model_manager.initialize()
        
        # Initialize data manager
        self.data_manager.initialize()
        
        # Register all agents
        self.register_agents()
        
        self.is_initialized = True
        logger.info("MediAssist Agent System initialized successfully")
        return True
        
    def register_agents(self):
        """Register all available specialized agents"""
        self.agents = {
            "router": RouterAgent(self),
            "clinical": ClinicalAgent(self),
            "literature": LiteratureAgent(self),
            "symptom": SymptomAgent(self),
            "drug": DrugAgent(self),
            "diet": DietAgent(self),
            "image": ImageAgent(self),
            "search": SearchAgent(self),
            "reflection": ReflectionAgent(self),
            "report": ReportAnalysisAgent(self)  # New agent for report analysis
        }
        logger.info(f"Registered {len(self.agents)} agents")
        
    def process_query(self, query_text: str, query_type: Optional[str] = None, 
                     additional_data: Optional[Dict] = None) -> Dict:
        """Process a user query and route to appropriate agents"""
        try:
            logger.info(f"Processing query: '{query_text[:50]}...' (Type: {query_type})")
            
            # Check initialization
            if not self.is_initialized:
                return {
                    "status": "error",
                    "response": "Agent system is not initialized yet. Please try again in a moment."
                }
            
            # Create a task object
            task = {
                "query": query_text,
                "type": query_type,
                "additional_data": additional_data or {},
                "status": "pending",
                "start_time": time.time()
            }
            
            # If query type is specified, route directly to that agent
            if query_type and query_type in self.agents:
                response = self.agents[query_type].process(task)
            else:
                # Otherwise use the router agent to determine the best agent
                response = self.agents["router"].process(task)
            
            # Always pass through reflection agent for quality check
            final_response = self.agents["reflection"].process({
                "original_query": query_text,
                "agent_response": response,
                "agent_type": query_type or "auto"
            })
            
            # Calculate processing time
            processing_time = time.time() - task["start_time"]
            final_response["processing_time"] = f"{processing_time:.2f}s"
            
            return final_response
            
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            return {
                "status": "error",
                "response": f"An error occurred while processing your query: {str(e)}"
            }
    
    def process_image_analysis(self, image_data: bytes, prompt: str) -> Dict:
        """Process an image analysis task"""
        try:
            logger.info(f"Processing image analysis with prompt: '{prompt[:50]}...'")
            
            # Check initialization
            if not self.is_initialized:
                return {
                    "status": "error",
                    "response": "Agent system is not initialized yet. Please try again in a moment."
                }
            
            # Create a task object
            task = {
                "image_data": image_data,
                "prompt": prompt,
                "status": "pending",
                "start_time": time.time()
            }
            
            # Route to image agent
            response = self.agents["image"].process(task)
            
            # Calculate processing time
            processing_time = time.time() - task["start_time"]
            response["processing_time"] = f"{processing_time:.2f}s"
            
            return response
            
        except Exception as e:
            logger.error(f"Error processing image analysis: {str(e)}")
            return {
                "status": "error",
                "response": f"An error occurred while analyzing your image: {str(e)}"
            }
    
    def process_report_analysis(self, reports_data: List[Dict], patient_info: Dict = None) -> Dict:
        """Process medical report analysis task"""
        try:
            logger.info(f"Processing report analysis for {len(reports_data)} reports")
            
            # Check initialization
            if not self.is_initialized:
                return {
                    "status": "error",
                    "response": "Agent system is not initialized yet. Please try again in a moment."
                }
            
            # Create a task object
            task = {
                "reports_data": reports_data,
                "patient_info": patient_info or {},
                "status": "pending",
                "start_time": time.time()
            }
            
            # Route to report analysis agent
            response = self.agents["report"].process(task)
            
            # Calculate processing time
            processing_time = time.time() - task["start_time"]
            response["processing_time"] = f"{processing_time:.2f}s"
            
            return response
            
        except Exception as e:
            logger.error(f"Error processing report analysis: {str(e)}")
            return {
                "status": "error",
                "response": f"An error occurred while analyzing your reports: {str(e)}"
            }
    
    def process_diet_plan(self, user_data: Dict) -> Dict:
        """Process a diet plan generation task"""
        try:
            logger.info(f"Processing diet plan for user profile")
            
            # Check initialization
            if not self.is_initialized:
                return {
                    "status": "error",
                    "response": "Agent system is not initialized yet. Please try again in a moment."
                }
            
            # Create a task object
            task = {
                "user_data": user_data,
                "status": "pending",
                "start_time": time.time()
            }
            
            # Route to diet agent
            response = self.agents["diet"].process(task)
            
            # Calculate processing time
            processing_time = time.time() - task["start_time"]
            response["processing_time"] = f"{processing_time:.2f}s"
            
            return response
            
        except Exception as e:
            logger.error(f"Error generating diet plan: {str(e)}")
            return {
                "status": "error",
                "response": f"An error occurred while generating your diet plan: {str(e)}"
            }
    
    def get_embedding(self, text: str) -> Optional[List[float]]:
        """Get embedding for text with caching"""
        # Check cache first
        cache_key = hash(text)
        if cache_key in self.embedding_cache:
            return self.embedding_cache[cache_key]
        
        # Not in cache, generate new embedding
        embedding = self.model_manager.get_embedding(text)
        
        # Cache it if valid
        if embedding:
            self.embedding_cache[cache_key] = embedding
            
        return embedding
        
    def refresh_embeddings(self, category: str) -> Dict:
        """Refresh embeddings for a specific category"""
        try:
            logger.info(f"Refreshing embeddings for {category} category")
            
            # Check if category exists
            if not self.data_manager.category_exists(category):
                return {
                    "status": "error",
                    "message": f"Invalid category: {category}"
                }
                
            # Refresh embeddings
            start_time = time.time()
            result = self.data_manager.refresh_embeddings(category)
            
            if not result["success"]:
                return {
                    "status": "error",
                    "message": result["message"]
                }
                
            processing_time = time.time() - start_time
            
            return {
                "status": "success",
                "message": f"Embeddings for {category} refreshed successfully",
                "processing_time": f"{processing_time:.2f}s",
                "embedding_count": result["count"]
            }
            
        except Exception as e:
            logger.error(f"Error refreshing embeddings: {str(e)}")
            return {
                "status": "error",
                "message": f"Failed to refresh embeddings: {str(e)}"
            }
    
    def generate_specialists(self, symptoms: str) -> Dict:
        """Generate specialist recommendations based on symptoms"""
        try:
            logger.info(f"Generating specialist recommendations for: '{symptoms[:50]}...'")
            
            # Generate recommendations using Azure OpenAI
            specialist_prompt = f"""Based on the following symptoms, generate recommendations for 5 medical specialists in Hyderabad, India who would be most appropriate to consult. For each specialist, provide their name, specialty, hospital/clinic affiliation (must be a hospital or clinic in Hyderabad, India), years of experience, and a brief expert bio (1-2 sentences). Use Indian names for the doctors.

Symptoms: {symptoms}

Format the response as a structured JSON array of 5 specialists with the following fields:
- id (integer from 1-5)
- name (doctor's full name with title, use Indian names)
- specialty (their medical specialty)
- hospital (where they practice in Hyderabad, India - use real hospital names)
- experience (years of experience as a number)
- bio (brief professional description)
- image_url (leave as "https://randomuser.me/api/portraits/men/[ID].jpg" where [ID] is between 1-99)

Assign each doctor a unique ID (1-5), and ensure the doctors are truly appropriate for the symptoms described, with relevant specialties. The response should ONLY contain the JSON array, no other text.
"""
            
            # Generate response using Azure OpenAI
            response = self.model_manager.generate_response(
                prompt=specialist_prompt,
                max_tokens=2048,
                temperature=0.2
            )
            
            # Parse the response to extract the JSON
            try:
                doctors_text = response
                # In case there's markdown formatting
                if "```json" in doctors_text:
                    doctors_text = doctors_text.split("```json")[1].split("```")[0].strip()
                elif "```" in doctors_text:
                    doctors_text = doctors_text.split("```")[1].split("```")[0].strip()
                
                doctors = json.loads(doctors_text)
                
                # Generate time slots for each doctor for the next 7 days
                doctors_with_slots = self.generate_time_slots_for_doctors(doctors)
                
                return {
                    "status": "success",
                    "specialists": doctors_with_slots
                }
            except Exception as e:
                logger.error(f"Error parsing specialist response: {str(e)}")
                return {
                    "status": "error",
                    "message": f"Failed to parse specialist recommendations: {str(e)}"
                }
                
        except Exception as e:
            logger.error(f"Error generating specialists: {str(e)}")
            return {
                "status": "error",
                "message": f"Failed to generate specialist recommendations: {str(e)}"
            }
    
    def generate_time_slots_for_doctors(self, doctors: List[Dict]) -> List[Dict]:
        """Generate time slots for each doctor"""
        for doctor in doctors:
            time_slots = []
            doctor_id = doctor["id"]
            
            # Generate slots for the next 7 days
            for day in range(7):
                date = datetime.now() + timedelta(days=day)
                # Create slots from 9AM to 4PM
                for hour in range(9, 16):
                    start_time = datetime(date.year, date.month, date.day, hour, 0)
                    end_time = start_time + timedelta(minutes=45)
                    time_slots.append({
                        "doctor_id": doctor_id,
                        "start_time": start_time,
                        "end_time": end_time,
                        "formatted_time": start_time.strftime("%A, %B %d at %I:%M %p"),
                        "is_booked": False
                    })
            
            doctor["available_slots"] = time_slots
        
        return doctors


# Base Agent class
class Agent:
    def __init__(self, system):
        self.system = system
        self.name = "base_agent"
        
    def process(self, task: Dict) -> Dict:
        """Process a task and return results"""
        raise NotImplementedError("Each agent must implement its own process method")
    
    def _create_base_response(self, status="success") -> Dict:
        """Create a base response template"""
        return {
            "status": status,
            "response": "",
            "agent": self.name,
            "timestamp": time.time()
        }
        
    def _generate_direct_response(self, query, category):
        """Generate a direct response using Azure OpenAI when no similar records are found"""
        try:
            # Create an effective prompt for the category
            prompt_templates = {
                'clinical': f"""As a medical AI assistant, provide a detailed clinical analysis for this query:

Query: {query}

Provide a comprehensive clinical analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🏥 Clinical Analysis

## 💊 Treatment Considerations

## ⚠️ Potential Complications

## 📊 Expected Outcomes

## 📋 Recommendations

For each section, provide detailed and evidence-based medical analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms.""",

                'literature': f"""As a medical research assistant, provide a literature-based analysis for this query:

Query: {query}

Provide a comprehensive research perspective with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 📚 Research Overview

## 🔬 Key Scientific Findings

## 📈 Evidence Summary

## 📊 Clinical Implications

## 🔮 Future Research Directions

For each section, provide detailed analysis based on current medical literature. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms.""",

                'symptom': f"""As a diagnostic assistant, analyze these symptoms:

Query: {query}

Provide a symptom analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🔍 Possible Causes

## ⚠️ Important Considerations

## 👨‍⚕️ Medical Advice

## 🚨 When to Seek Help

## 📋 Next Steps

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms.""",

                'drug': f"""As a pharmaceutical expert, provide information about this medication query:

Query: {query}

Provide a comprehensive medication analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 💊 Medication Information

## ⚠️ Important Considerations

## 👁️ Side Effects

## 🔄 Interactions

## 📋 Usage Guidelines

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms.""",

                'diet': f"""As a nutrition expert, provide dietary advice for this query:

Query: {query}

Provide comprehensive nutritional guidance with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🥗 Nutritional Analysis

## ⚖️ Dietary Recommendations

## 🍎 Food Suggestions

## 💧 Hydration Guidelines

## 🏋️ Lifestyle Considerations

For each section, provide detailed and practical advice. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms.""",

                'report': f"""As a medical report analysis expert, provide comprehensive analysis for this report data:

Query: {query}

Provide detailed medical report analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 📋 Report Summary

## 🔍 Key Findings

## ⚠️ Abnormal Values

## 💊 Clinical Significance

## 📊 Recommendations

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms."""
            }
            
            # Select the appropriate prompt template
            prompt = prompt_templates.get(category, prompt_templates['clinical'])
            
            # Generate response using Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            return response
            
        except Exception as e:
            logger.error(f"Error generating direct response: {str(e)}")
            return f"I apologize, but I encountered an error while processing your query. Please try rephrasing your question or providing more details."


# Report Analysis Agent - New agent for analyzing medical reports
class ReportAnalysisAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "report"
        
    def process(self, task: Dict) -> Dict:
        """Process medical report analysis tasks"""
        try:
            reports_data = task["reports_data"]
            patient_info = task.get("patient_info", {})
            
            # Combine all report texts
            combined_text = ""
            report_summaries = []
            
            for report in reports_data:
                combined_text += f"\n--- {report['filename']} ---\n{report['content']}\n"
                report_summaries.append(f"- {report['filename']} ({report['size']} bytes)")
            
            # Extract health values from all reports
            health_values = extract_health_values(combined_text)
            
            # Generate health charts if values found
            chart_data = None
            if health_values:
                gender = patient_info.get('gender', 'unknown')
                chart_data = generate_health_chart(health_values, gender)
            
            # Create comprehensive analysis prompt
            analysis_prompt = f"""As a medical report analysis expert, provide a comprehensive analysis of these medical reports.

Patient Information:
{json.dumps(patient_info, indent=2) if patient_info else "Not provided"}

Reports Analyzed:
{chr(10).join(report_summaries)}

Combined Report Content:
{combined_text[:8000]}  # Limit to prevent token overflow

Extracted Health Values:
{json.dumps(health_values, indent=2) if health_values else "None automatically extracted"}

Provide a detailed medical report analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 📋 Executive Summary

## 🔍 Key Findings Analysis

## ⚠️ Abnormal Values & Concerns

## 📈 Trend Analysis (if multiple reports)

## 💊 Clinical Significance

## 🩺 Recommended Actions

## 📊 Follow-up Recommendations

## ⚕️ Specialist Consultations

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. 

If specific health values are found, compare them with normal ranges and explain their clinical significance. If trends can be identified across multiple reports, highlight them.

Focus on:
1. Critical findings that need immediate attention
2. Values outside normal ranges
3. Patterns or trends
4. Recommendations for patient care
5. Suggested follow-up tests or consultations"""
            
            # Generate response using Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=analysis_prompt,
                max_tokens=4096,
                temperature=0.3
            )
            
            # Create response object
            response_obj = self._create_base_response()
            response_obj["response"] = response
            response_obj["health_values"] = health_values
            response_obj["chart_data"] = chart_data
            response_obj["reports_processed"] = len(reports_data)
            
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in report analysis agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error analyzing your medical reports: {str(e)}"
            return response


# Router Agent - Routes queries to the appropriate specialized agent
class RouterAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "router"
        
    def process(self, task: Dict) -> Dict:
        """Determine the best agent to handle this query"""
        try:
            query = task["query"]
            
            # Force symptom agent for symptom-like queries
            if any(word in query.lower() for word in ["symptom", "diagnose", "suffering from", "experiencing", "feeling", "pain", "ache", "headache", "fever", "cough", "rash"]):
                agent_name = "symptom"
                logger.info(f"Router detected symptom keywords, using symptom agent for query: '{query[:50]}...'")
                
                # Create a new task with the determined agent type
                agent_task = task.copy()
                agent_task["type"] = agent_name
                
                # Route to the selected agent
                response = self.system.agents[agent_name].process(agent_task)
                
                # Add the determined agent name to the response
                response["agent"] = agent_name
                
                return response
            
            # Generate router prompt
            router_prompt = f"""As an AI medical assistant router, analyze this query and determine which specialized medical agent should handle it.

Query: {query}

Available agents:
1. "clinical" - For clinical cases, patient treatment, medical advice
2. "literature" - For medical research, studies, papers, evidence-based medicine
3. "symptom" - For symptom analysis, diagnoses, medical conditions
4. "drug" - For medication information, drug interactions, pharmaceutical queries
5. "diet" - For nutrition, diet plans, food-related health queries
6. "report" - For medical report analysis, lab results, health data interpretation

Respond with ONLY ONE word - the name of the most appropriate agent (clinical, literature, symptom, drug, diet, report). Choose the single best match.
"""
            
            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=router_prompt,
                max_tokens=10,
                temperature=0.1
            )
            agent_name = response.strip().lower()
            
            # Validate the agent name
            valid_agents = ["clinical", "literature", "symptom", "drug", "diet", "report"]
            if agent_name not in valid_agents:
                logger.warning(f"Router returned invalid agent: {agent_name}, defaulting to clinical")
                agent_name = "clinical"
                
            logger.info(f"Router selected {agent_name} agent for query: '{query[:50]}...'")
            
            # Create a new task with the determined agent type
            agent_task = task.copy()
            agent_task["type"] = agent_name
            
            # Route to the selected agent
            response = self.system.agents[agent_name].process(agent_task)
            
            # Add the determined agent name to the response
            response["agent"] = agent_name
            
            return response
            
        except Exception as e:
            logger.error(f"Error in router agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error determining how to process your query. Please try again."
            return response


# Clinical Agent - Handles clinical cases and medical advice
class ClinicalAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "clinical"
        
    def process(self, task: Dict) -> Dict:
        """Process clinical queries"""
        try:
            query = task["query"]
            
            # Find similar clinical cases
            similar_cases = self.system.data_manager.find_similar(query, "clinical", top_k=3)
            
            if not similar_cases or len(similar_cases) == 0:
                logger.info(f"No similar clinical cases found for query: '{query[:50]}...', generating direct response")
                
                # Generate a direct response using Azure OpenAI
                direct_response = self._generate_direct_response(query, "clinical")
                
                response = self._create_base_response()
                response["response"] = direct_response
                response["direct_response"] = True  # Flag to indicate this was a direct response
                return response
            
            # Prepare clinical prompt with similar cases
            formatted_cases = self._format_similar_cases(similar_cases)
            
            clinical_prompt = f"""As a clinical decision support agent, analyze this medical query based on similar cases in our database.

User Query: {query}

Similar Clinical Cases:
{formatted_cases}

Provide a comprehensive clinical analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🏥 Case Similarity Analysis

## 💊 Evidence-Based Treatment Recommendations

## ⚠️ Potential Complications to Monitor

## 📊 Expected Outcomes

## 📋 Follow-up Recommendations

For each section, provide detailed medical analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative."""

            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=clinical_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            response_obj = self._create_base_response()
            response_obj["response"] = response
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in clinical agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error processing your clinical query. Please try again."
            return response
    
    def _format_similar_cases(self, cases: List[Dict]) -> str:
        """Format similar cases for the prompt"""
        formatted = ""
        for i, case in enumerate(cases):
            record = case["record"]
            formatted += f"Case {i+1} (Similarity: {case['similarity']:.2f}):\n"
            formatted += f"- Age: {record.get('age', 'N/A')}, Gender: {record.get('gender', 'N/A')}\n"
            formatted += f"- Symptoms: {record.get('symptoms', 'N/A')}\n"
            formatted += f"- Diagnosis: {record.get('diagnosis', 'N/A')}\n"
            formatted += f"- Treatment: {record.get('treatment', 'N/A')}\n"
            formatted += f"- Outcome: {record.get('outcome', 'N/A')}\n\n"
        return formatted


# Literature Agent - Handles medical research and studies
class LiteratureAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "literature"
        
    def process(self, task: Dict) -> Dict:
        """Process literature queries"""
        try:
            query = task["query"]
            
            # Find similar literature
            similar_literature = self.system.data_manager.find_similar(query, "literature", top_k=3)
            
            if not similar_literature or len(similar_literature) == 0:
                logger.info(f"No similar literature found for query: '{query[:50]}...', generating direct response")
                
                # Generate a direct response using Azure OpenAI
                direct_response = self._generate_direct_response(query, "literature")
                
                response = self._create_base_response()
                response["response"] = direct_response
                response["direct_response"] = True
                return response
            
            # Prepare literature prompt with similar papers
            formatted_literature = self._format_similar_literature(similar_literature)
            
            literature_prompt = f"""As a medical research assistant, analyze this research query based on our literature database.

User Query: {query}

Relevant Literature:
{formatted_literature}

Provide a comprehensive literature review with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 📚 Relevant Studies Analysis

## 🔬 Key Findings Synthesis

## 📈 Treatment Efficacy Data

## 📊 Statistical Evidence

## 🔮 Research Gaps & Future Directions

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative."""

            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=literature_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            response_obj = self._create_base_response()
            response_obj["response"] = response
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in literature agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error processing your literature query. Please try again."
            return response
    
    def _format_similar_literature(self, literature: List[Dict]) -> str:
        """Format similar literature for the prompt"""
        formatted = ""
        for i, paper in enumerate(literature):
            record = paper["record"]
            formatted += f"Paper {i+1} (Similarity: {paper['similarity']:.2f}):\n"
            formatted += f"- Title: {record.get('title', 'N/A')}\n"
            formatted += f"- Authors: {record.get('authors', 'N/A')}\n"
            formatted += f"- Journal: {record.get('journal', 'N/A')}, Date: {record.get('publication_date', 'N/A')}\n"
            formatted += f"- Key Findings: {record.get('key_findings', 'N/A')}\n"
            formatted += f"- Methodology: {record.get('methodology', 'N/A')}\n\n"
        return formatted


# Symptom Agent - Handles symptom analysis and diagnostic suggestions
class SymptomAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "symptom"
        
    def process(self, task: Dict) -> Dict:
        """Process symptom queries"""
        try:
            query = task["query"]
            
            # Find similar symptom cases
            similar_symptoms = self.system.data_manager.find_similar(query, "symptom", top_k=3)
            
            if not similar_symptoms or len(similar_symptoms) == 0:
                logger.info(f"No similar symptom information found for query: '{query[:50]}...', generating direct response")
                
                # Generate a direct response using Azure OpenAI
                direct_response = self._generate_direct_response(query, "symptom")
                
                # Get specialist recommendations based on the symptoms
                specialists_result = self.system.generate_specialists(query)
                
                response = self._create_base_response()
                response["response"] = direct_response
                response["direct_response"] = True
                # Add the flag to show booking interface
                response["show_booking"] = True
                # Add specialists data
                if specialists_result["status"] == "success":
                    response["specialists"] = specialists_result["specialists"]
                
                return response
            
            # Prepare symptom prompt with similar cases
            formatted_symptoms = self._format_similar_symptoms(similar_symptoms)
            
            symptom_prompt = f"""As a diagnostic assistant, analyze these symptoms based on our symptom database.

User Query: {query}

Relevant Symptom Cases:
{formatted_symptoms}

Provide a symptom analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🔍 Potential Diagnoses

## ⚠️ Key Risk Factors

## 👨‍⚕️ Specialist Recommendations

## 🚨 Urgency Assessment

## 📋 Recommended Diagnostic Tests

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative."""

            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=symptom_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            # Get specialist recommendations based on the symptoms
            specialists_result = self.system.generate_specialists(query)
            
            response_obj = self._create_base_response()
            response_obj["response"] = response
            # Add the flag to show booking interface
            response_obj["show_booking"] = True
            # Add specialists data
            if specialists_result["status"] == "success":
                response_obj["specialists"] = specialists_result["specialists"]
            
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in symptom agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error processing your symptom query. Please try again."
            return response
    
    def _format_similar_symptoms(self, symptoms: List[Dict]) -> str:
        """Format similar symptom cases for the prompt"""
        formatted = ""
        for i, symptom in enumerate(symptoms):
            record = symptom["record"]
            formatted += f"Case {i+1} (Similarity: {symptom['similarity']:.2f}):\n"
            formatted += f"- Presenting Symptoms: {record.get('presenting_symptoms', 'N/A')}\n"
            formatted += f"- Diagnosis: {record.get('diagnosis', 'N/A')}\n"
            formatted += f"- Risk Factors: {record.get('risk_factors', 'N/A')}\n"
            formatted += f"- Urgency Level: {record.get('urgency_level', 'N/A')}\n"
            formatted += f"- Recommended Tests: {record.get('diagnostic_tests', 'N/A')}\n\n"
        return formatted


# Drug Agent - Handles medication information and interactions
class DrugAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "drug"
        
    def process(self, task: Dict) -> Dict:
        """Process drug queries"""
        try:
            query = task["query"]
            
            # Find similar drug interactions
            similar_drugs = self.system.data_manager.find_similar(query, "drug", top_k=3)
            
            if not similar_drugs or len(similar_drugs) == 0:
                logger.info(f"No similar medication information found for query: '{query[:50]}...', generating direct response")
                
                # Generate a direct response using Azure OpenAI
                direct_response = self._generate_direct_response(query, "drug")
                
                response = self._create_base_response()
                response["response"] = direct_response
                response["direct_response"] = True
                return response
            
            # Prepare drug prompt with similar interactions
            formatted_drugs = self._format_similar_drugs(similar_drugs)
            
            drug_prompt = f"""As a pharmaceutical expert, analyze these medication interactions.

User Query: {query}

Relevant Drug Interactions:
{formatted_drugs}

Provide a comprehensive interaction analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## ⚠️ Interaction Severity Assessment

## 👁️ Effects to Monitor

## 💊 Medication Adjustments

## 🔄 Alternative Medications

## 📋 Patient Monitoring Guidelines

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative."""

            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=drug_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            response_obj = self._create_base_response()
            response_obj["response"] = response
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in drug agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error processing your medication query. Please try again."
            return response
    
    def _format_similar_drugs(self, drugs: List[Dict]) -> str:
        """Format similar drug interactions for the prompt"""
        formatted = ""
        for i, drug in enumerate(drugs):
            record = drug["record"]
            formatted += f"Interaction {i+1} (Similarity: {drug['similarity']:.2f}):\n"
            formatted += f"- Medications: {record.get('medications', 'N/A')}\n"
            formatted += f"- Severity: {record.get('severity', 'N/A')}\n"
            formatted += f"- Effects: {record.get('effects', 'N/A')}\n"
            formatted += f"- Recommendations: {record.get('recommendations', 'N/A')}\n"
            formatted += f"- Alternatives: {record.get('alternatives', 'N/A')}\n\n"
        return formatted


# Diet Agent - Handles nutrition and diet plan generation
class DietAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "diet"
        
    def process(self, task: Dict) -> Dict:
        """Process diet queries or generate diet plans"""
        try:
            # Check if this is a diet plan generation task
            if "user_data" in task:
                return self._generate_diet_plan(task["user_data"])
            
            # Otherwise, treat as a general diet query
            query = task["query"]
            
            # Find similar diet information
            similar_diets = self.system.data_manager.find_similar(query, "diet", top_k=3)
            
            if not similar_diets or len(similar_diets) == 0:
                logger.info(f"No similar nutrition information found for query: '{query[:50]}...', generating direct response")
                
                # Generate a direct response using Azure OpenAI
                direct_response = self._generate_direct_response(query, "diet")
                
                response = self._create_base_response()
                response["response"] = direct_response
                response["direct_response"] = True
                return response
            
            # Prepare diet prompt with similar information
            formatted_diets = self._format_similar_diets(similar_diets)
            
            diet_prompt = f"""As a nutritionist and diet planning expert, analyze this dietary query.

User Query: {query}

Relevant Diet Information:
{formatted_diets}

Provide a comprehensive dietary analysis with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🥗 Nutritional Assessment

## ⚖️ Dietary Recommendations

## 🍎 Food Suggestions

## 💧 Hydration Guidelines

## 🏋️ Exercise Recommendations

## ⚠️ Dietary Considerations

For each section, provide detailed analysis. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative."""

            # Get response from Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=diet_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            response_obj = self._create_base_response()
            response_obj["response"] = response
            return response_obj
            
        except Exception as e:
            logger.error(f"Error in diet agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error processing your nutrition query. Please try again."
            return response
    
    def _format_similar_diets(self, diets: List[Dict]) -> str:
        """Format similar diet information for the prompt"""
        formatted = ""
        for i, diet in enumerate(diets):
            record = diet["record"]
            formatted += f"Diet Plan {i+1} (Similarity: {diet['similarity']:.2f}):\n"
            formatted += f"- Age Group: {record.get('age_group', 'N/A')}\n"
            formatted += f"- Weight Range: {record.get('weight_range', 'N/A')}\n"
            formatted += f"- Height Range: {record.get('height_range', 'N/A')}\n"
            formatted += f"- Health Goals: {record.get('health_goals', 'N/A')}\n"
            formatted += f"- Dietary Preferences: {record.get('dietary_preferences', 'N/A')}\n"
            formatted += f"- Recommended Foods: {record.get('recommended_foods', 'N/A')}\n\n"
        return formatted
    
    def _generate_diet_plan(self, user_data: Dict) -> Dict:
        """Generate a personalized diet plan based on user data"""
        try:
            # Extract user data
            age = user_data.get('age', '')
            gender = user_data.get('gender', '')
            height = user_data.get('height', '')
            weight = user_data.get('weight', '')
            health_goal = user_data.get('health_goal', '')
            activity_level = user_data.get('activity_level', '')
            dietary_preferences = user_data.get('dietary_preferences', '')
            medical_conditions = user_data.get('medical_conditions', '')
            allergies = user_data.get('allergies', '')
            supplements = user_data.get('supplements', '')
            
            # Create diet plan prompt
            diet_plan_prompt = f"""As a professional nutritionist and dietitian, create a personalized 7-day diet plan for this individual.

User Profile:
- Age: {age}
- Gender: {gender}
- Height: {height}
- Weight: {weight}
- Health Goal: {health_goal}
- Activity Level: {activity_level}
- Dietary Preferences: {dietary_preferences}
- Medical Conditions: {medical_conditions}
- Allergies/Intolerances: {allergies}
- Supplements: {supplements}

Provide a comprehensive nutritional plan with CLEAN markdown formatting. Each section should start with the exact headings below with emojis:

## 🥗 Nutritional Assessment

## ⚖️ Caloric and Macronutrient Recommendations

## 💧 Hydration Guidelines

## 7-Day Meal Plan
### 🍎 Day 1
### 🍎 Day 2
### 🍎 Day 3
### 🍎 Day 4
### 🍎 Day 5
### 🍎 Day 6
### 🍎 Day 7

## 🏋️ Exercise Recommendations

## 📋 Grocery List

## ⚠️ Special Considerations

For each section, provide detailed, practical guidance tailored to this individual. Format your response with bullet points using * (not -), use numbered lists (1. 2.) where appropriate, and use **bold text** for important information or terms. Keep your response concise but informative.

Provide SPECIFIC MEAL IDEAS for each day that are practical, accessible, and aligned with the dietary preferences specified. Include approximate portion sizes where appropriate.

For the grocery list, organize items by category (proteins, fruits, vegetables, grains, etc.) to make shopping easier.
"""
            
            # Generate response using Azure OpenAI
            response = self.system.model_manager.generate_response(
                prompt=diet_plan_prompt,
                max_tokens=4096,
                temperature=0.4
            )
            
            result = self._create_base_response()
            result["response"] = response
            return result
            
        except Exception as e:
            logger.error(f"Error generating diet plan: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error generating your diet plan. Please try again."
            return response


# Image Agent - Handles medical image analysis
class ImageAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "image"
        
    def process(self, task: Dict) -> Dict:
        """Process image analysis tasks"""
        try:
            image_data = task["image_data"]
            prompt = task["prompt"]
            
            # Create a unique key for caching
            import hashlib
            img_hash = hashlib.md5(image_data).hexdigest()
            prompt_hash = hashlib.md5(prompt.encode('utf-8')).hexdigest()
            cache_key = f"{img_hash}_{prompt_hash}"
            
            # Check if we have a cached result
            if cache_key in self.system.results_cache:
                logger.info(f"Using cached image analysis result for {cache_key}")
                cached_result = self.system.results_cache[cache_key]
                response = self._create_base_response()
                response["response"] = cached_result
                response["cached"] = True
                return response
            
            # Process the image
            image = Image.open(io.BytesIO(image_data))
            
            # Enhanced medical image prompt
            enhanced_prompt = f"""As a medical image analysis expert, analyze this medical image with precision and clinical relevance.

Medical Context: {prompt}

Please provide your analysis in a clear, structured format with the following sections:
1. Image Description - Describe what you see in the image
2. Key Findings - Identify notable features or abnormalities
3. Possible Interpretations - Discuss what these findings might indicate
4. Recommendations - Suggest next steps or further tests if applicable

Format your response with clear markdown headings and bullet points for readability."""
            
            # Try Azure OpenAI first, then fallback to Google AI
            response_text = None
            
            # Try Azure OpenAI with GPT-4 Vision capabilities
            try:
                # Convert image to base64
                import base64
                from io import BytesIO
                
                buffer = BytesIO()
                image.save(buffer, format="PNG")
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                # Use Azure OpenAI for image analysis with gpt-4o (higher rate limits)
                azure_response = self.system.model_manager.default_model.chat.completions.create(
                    model="gpt-4o",  # Use gpt-4o instead of gpt-4o-mini for image analysis
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": enhanced_prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{img_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4096,
                    temperature=0.4
                )
                
                response_text = azure_response.choices[0].message.content
                logger.info("Successfully used Azure OpenAI for image analysis")
                
            except Exception as azure_error:
                logger.warning(f"Azure OpenAI image analysis failed: {str(azure_error)}")
                
                # Fallback to Google AI for image analysis
                if google_ai_model:
                    try:
                        # Use Google AI as fallback
                        google_response = google_ai_model.generate_content([enhanced_prompt, image])
                        response_text = google_response.text
                        logger.info("Successfully used fallback image analysis")
                    except Exception as google_error:
                        logger.error(f"Google AI fallback also failed: {str(google_error)}")
                        response_text = "I apologize, but I'm currently unable to analyze images. Please try again later or contact support."
                else:
                    response_text = "I apologize, but image analysis is currently unavailable. Please try again later."
            
            # Cache the result
            if response_text:
                self.system.results_cache[cache_key] = response_text
            
            # Format and return the response
            response = self._create_base_response()
            response["response"] = response_text
            return response
            
        except Exception as e:
            logger.error(f"Error in image agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error analyzing your medical image. Please try again."
            return response


# Search Agent - Handles external searches for supplementary information
class SearchAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "search"
        
    def process(self, task: Dict) -> Dict:
        """Process search tasks - simulated as we don't have actual search integration"""
        try:
            query = task.get("query", "")
            
            # This is a simulation as we don't have actual search integration
            response = self._create_base_response()
            response["response"] = f"I've searched for additional information on '{query}', but external search capabilities are not currently enabled. I'll provide information based on my existing knowledge."
            
            return response
            
        except Exception as e:
            logger.error(f"Error in search agent: {str(e)}")
            response = self._create_base_response("error")
            response["response"] = f"Error performing search. Please try again."
            return response


# Reflection Agent - Checks output quality and accuracy
class ReflectionAgent(Agent):
    def __init__(self, system):
        super().__init__(system)
        self.name = "reflection"
        
    def process(self, task: Dict) -> Dict:
        """Process reflection tasks - review other agent outputs"""
        try:
            original_query = task.get("original_query", "")
            agent_response = task.get("agent_response", {})
            agent_type = task.get("agent_type", "unknown")
            
            # If the response is already an error, just pass it through
            if agent_response.get("status") != "success":
                return agent_response
            
            # Check the response text
            response_text = agent_response.get("response", "")
            
            # Check if response appears too generic or unhelpful
            if len(response_text) < 100:
                logger.warning(f"Response from {agent_type} agent seems too short, adding disclaimer")
                response_text += "\n\n*Note: This response is limited. For more comprehensive information, please provide additional details or consult a healthcare professional.*"
            
            # Check if response contains medical disclaimer if needed
            if any(keyword in original_query.lower() for keyword in ["diagnosis", "treatment", "cure", "remedy"]):
                if "professional" not in response_text.lower() and "consult" not in response_text.lower():
                    response_text += "\n\n*Disclaimer: This information is for educational purposes only and is not a substitute for professional medical advice. Always consult with a qualified healthcare provider for medical concerns.*"
            
            # Return the potentially modified response
            agent_response["response"] = response_text
            return agent_response
            
        except Exception as e:
            logger.error(f"Error in reflection agent: {str(e)}")
            # If there's an error in reflection, return the original response
            return task.get("agent_response", {
                "status": "error",
                "response": "Error processing your request."
            })


# Data Manager - Handles data loading, embedding, and retrieval
class DataManager:
    def __init__(self, model_manager=None):
        self.data = {}
        self.embeddings = {}
        self.data_dir = "data"
        self.embeddings_dir = "data/embeddings"
        self.data_lock = threading.RLock()
        self.embedding_lock = threading.RLock()
        self.model_manager = model_manager
        
    def initialize(self):
        """Initialize the data manager"""
        logger.info("Initializing Data Manager")
        
        # Create necessary directories
        os.makedirs(self.data_dir, exist_ok=True)
        os.makedirs(self.embeddings_dir, exist_ok=True)
        
        # Load data
        APP_STATE["loading_progress"]["data"] = {"status": "in_progress", "message": "Loading data files"}
        self.load_data()
        
        # Load embeddings
        APP_STATE["loading_progress"]["embeddings"] = {"status": "in_progress", "message": "Loading embeddings"}
        self.load_embeddings()
        
        logger.info("Data Manager initialized successfully")
        return True
    
    def load_data(self):
        """Load all medical data from Excel files"""
        try:
            categories = [
                ('clinical', 'clinical_cases.xlsx'),
                ('literature', 'medical_literature.xlsx'),
                ('symptom', 'symptom_cases.xlsx'),
                ('drug', 'drug_interactions.xlsx'),
                ('diet', 'diet_plans.xlsx')
            ]
            
            success = True
            for category, filename in categories:
                try:
                    df = pd.read_excel(f'{self.data_dir}/{filename}')
                    self.data[category] = df.to_dict(orient='records')
                    logger.info(f"Loaded {len(self.data[category])} {category} records")
                except Exception as e:
                    logger.error(f"Error loading {category} data: {str(e)}")
                    success = False
            
            if success:
                APP_STATE["loading_progress"]["data"] = {"status": "complete", "message": f"Loaded {sum(len(v) for v in self.data.values())} total records"}
            else:
                APP_STATE["loading_progress"]["data"] = {"status": "warning", "message": "Some data files could not be loaded"}
                
        except Exception as e:
            logger.error(f"Error loading data: {str(e)}")
            APP_STATE["loading_progress"]["data"] = {"status": "error", "message": str(e)}
    
    def load_embeddings(self):
        """Load embeddings for all categories"""
        try:
            categories = ['clinical', 'literature', 'symptom', 'drug', 'diet']
            
            success = True
            for category in categories:
                try:
                    embedding_path = f"{self.embeddings_dir}/{category}_embeddings.pt"
                    if os.path.exists(embedding_path):
                        self.embeddings[category] = torch.load(embedding_path, map_location=torch.device('cpu'))
                        logger.info(f"Loaded {len(self.embeddings[category])} {category} embeddings")
                    else:
                        logger.info(f"No embeddings found for {category}, will generate when needed")
                except Exception as e:
                    logger.error(f"Error loading {category} embeddings: {str(e)}")
                    success = False
            
            if success:
                APP_STATE["loading_progress"]["embeddings"] = {"status": "complete", "message": "Embeddings loaded successfully"}
            else:
                APP_STATE["loading_progress"]["embeddings"] = {"status": "warning", "message": "Some embeddings could not be loaded"}
                
        except Exception as e:
            logger.error(f"Error loading embeddings: {str(e)}")
            APP_STATE["loading_progress"]["embeddings"] = {"status": "error", "message": str(e)}
    
    def category_exists(self, category: str) -> bool:
        """Check if a category exists in the data"""
        return category in self.data
    
    def find_similar(self, query: str, category: str, top_k: int = 3) -> List[Dict]:
        """Find records similar to the query in the specified category with improved handling"""
        try:
            # Ensure the category exists
            if category not in self.data:
                logger.warning(f"Category {category} not found in data")
                return []
            
            # Check if embeddings exist for this category
            with self.embedding_lock:
                if category not in self.embeddings or not self.embeddings[category]:
                    # Generate embeddings if they don't exist
                    logger.info(f"Generating embeddings for {category} category")
                    self.generate_embeddings(category)
                    
                    # Check if generation was successful
                    if category not in self.embeddings or not self.embeddings[category]:
                        logger.error(f"Failed to generate embeddings for {category}")
                        return []
            
            # Get embedding for the query from the shared model manager
            if not self.model_manager:
                logger.error("Model manager not available in DataManager")
                return []
                
            query_embedding = self.model_manager.get_embedding(query)
            
            if not query_embedding:
                logger.warning("Failed to get embedding for query")
                return []
            
            # Calculate similarities with error handling
            similarities = []
            query_embedding_np = np.array(query_embedding)
            
            for item in self.embeddings[category]:
                try:
                    # Skip items with missing or invalid embeddings
                    if 'embedding' not in item or not item['embedding']:
                        continue
                        
                    item_embedding = np.array(item['embedding'])
                    
                    # Skip zero vectors
                    if np.all(item_embedding == 0):
                        continue
                        
                    similarity = self.cosine_similarity(query_embedding_np, item_embedding)
                    
                    # Only include items with reasonable similarity 
                    # (lowered threshold to find more matches)
                    if similarity > 0.1:
                        similarities.append({
                            'record': item['record'],
                            'similarity': similarity
                        })
                except Exception as e:
                    logger.warning(f"Error calculating similarity for an item: {str(e)}")
                    continue
            
            # Sort by similarity and return top_k
            similarities.sort(key=lambda x: x['similarity'], reverse=True)
            
            # Log the number of similar records found
            logger.info(f"Found {len(similarities)} similar records for category {category}, returning top {min(top_k, len(similarities))}")
            
            return similarities[:top_k]
            
        except Exception as e:
            logger.error(f"Error finding similar records: {str(e)}")
            return []
    
    def generate_embeddings(self, category: str) -> bool:
        """Generate embeddings for a category"""
        try:
            # Check if category exists
            if category not in self.data:
                logger.warning(f"Category {category} not found in data")
                return False
            
            if not self.model_manager:
                logger.error("Model manager not available in DataManager")
                return False
                
            records = self.data[category]
            category_embeddings = []
            
            for record in records:
                # Create text representation
                text = self.prepare_text_for_embedding(record, category)
                
                # Get embedding
                embedding = self.model_manager.get_embedding(text)
                
                if embedding:
                    category_embeddings.append({
                        'record': record,
                        'embedding': embedding
                    })
                else:
                    logger.warning(f"Failed to get embedding for record in {category}")
            
            # Save embeddings
            with self.embedding_lock:
                self.embeddings[category] = category_embeddings
                
                # Save to file
                torch.save(category_embeddings, f"{self.embeddings_dir}/{category}_embeddings.pt")
            
            logger.info(f"Generated {len(category_embeddings)} embeddings for {category} category")
            return True
            
        except Exception as e:
            logger.error(f"Error generating embeddings for {category}: {str(e)}")
            return False
    
    def refresh_embeddings(self, category: str) -> Dict:
        """Refresh embeddings for a category"""
        try:
            # Check if category exists
            if category not in self.data:
                return {"success": False, "message": f"Category {category} not found in data", "count": 0}
            
            # Generate new embeddings
            success = self.generate_embeddings(category)
            
            if not success:
                return {"success": False, "message": f"Failed to generate embeddings for {category}", "count": 0}
            
            # Return result
            count = len(self.embeddings[category])
            return {"success": True, "message": f"Generated {count} embeddings for {category}", "count": count}
            
        except Exception as e:
            logger.error(f"Error refreshing embeddings for {category}: {str(e)}")
            return {"success": False, "message": str(e), "count": 0}
    
    def prepare_text_for_embedding(self, record: Dict, category: str) -> str:
        """Convert a record to text for embedding"""
        try:
            if category == 'clinical':
                return f"Case ID: {record.get('case_id', '')}. Patient: {record.get('age', '')} year old {record.get('gender', '')}. Symptoms: {record.get('symptoms', '')}. Medical history: {record.get('medical_history', '')}. Diagnosis: {record.get('diagnosis', '')}. Treatment: {record.get('treatment', '')}. Outcome: {record.get('outcome', '')}. Complications: {record.get('complications', '')}."
            
            elif category == 'literature':
                return f"Paper ID: {record.get('paper_id', '')}. Title: {record.get('title', '')}. Authors: {record.get('authors', '')}. Published: {record.get('publication_date', '')} in {record.get('journal', '')}. Key findings: {record.get('key_findings', '')}. Methodology: {record.get('methodology', '')}. Sample size: {record.get('sample_size', '')}."
            
            elif category == 'symptom':
                return f"Symptom ID: {record.get('symptom_id', '')}. Presenting symptoms: {record.get('presenting_symptoms', '')}. Diagnosis: {record.get('diagnosis', '')}. Risk factors: {record.get('risk_factors', '')}. Specialists: {record.get('recommended_specialists', '')}. Urgency: {record.get('urgency_level', '')}. Tests: {record.get('diagnostic_tests', '')}."
            
            elif category == 'drug':
                return f"Interaction ID: {record.get('interaction_id', '')}. Medications: {record.get('medications', '')}. Severity: {record.get('severity', '')}. Effects: {record.get('effects', '')}. Recommendations: {record.get('recommendations', '')}. Alternatives: {record.get('alternatives', '')}."
            
            elif category == 'diet':
                return f"Diet ID: {record.get('diet_id', '')}. Age Group: {record.get('age_group', '')}. Weight Range: {record.get('weight_range', '')}. Height Range: {record.get('height_range', '')}. Health Goals: {record.get('health_goals', '')}. Dietary Preferences: {record.get('dietary_preferences', '')}. Medical Conditions: {record.get('medical_conditions', '')}. Recommended Foods: {record.get('recommended_foods', '')}. Foods to Avoid: {record.get('foods_to_avoid', '')}. Exercise Recommendations: {record.get('exercise_recommendations', '')}."
            
            return ""
        except Exception as e:
            logger.error(f"Error preparing text for embedding: {str(e)}")
            return ""
    
    def cosine_similarity(self, v1: np.ndarray, v2: np.ndarray) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            v1 = v1.flatten()
            v2 = v2.flatten()
            
            dot_product = np.dot(v1, v2)
            norm_v1 = np.linalg.norm(v1)
            norm_v2 = np.linalg.norm(v2)
            
            if norm_v1 == 0 or norm_v2 == 0:
                return 0
                
            return float(dot_product / (norm_v1 * norm_v2))
        except Exception as e:
            logger.error(f"Error calculating cosine similarity: {str(e)}")
            return 0


# Model Manager - Handles LLM initialization and calls
class ModelManager:
    def __init__(self):
        self.models = {}
        self.default_model = None
        
    def initialize(self):
        """Initialize the model manager"""
        logger.info("Initializing Model Manager")
        
        APP_STATE["loading_progress"]["models"] = {"status": "in_progress", "message": "Initializing models"}
        
        try:
            # Initialize Azure OpenAI
            if not AZURE_OPENAI_API_KEY:
                logger.warning("AZURE_OPENAI_API_KEY not set")
                APP_STATE["loading_progress"]["models"] = {"status": "error", "message": "AZURE_OPENAI_API_KEY not set"}
                return False
                
            if not AZURE_OPENAI_ENDPOINT:
                logger.warning("AZURE_OPENAI_ENDPOINT not set")
                APP_STATE["loading_progress"]["models"] = {"status": "error", "message": "AZURE_OPENAI_ENDPOINT not set"}
                return False
                
            self.default_model = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_API_VERSION,
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                http_client=httpx.Client(verify=False)  # Disable SSL verification for corporate environments
            )
            logger.info(f"Initialized Azure OpenAI model: {AZURE_OPENAI_MODEL_NAME}")
            
            APP_STATE["loading_progress"]["models"] = {"status": "complete", "message": "Models initialized successfully"}
            return True
            
        except Exception as e:
            logger.error(f"Error initializing models: {str(e)}")
            APP_STATE["loading_progress"]["models"] = {"status": "error", "message": str(e)}
            return False
    
    def get_model(self, model_name: Optional[str] = None):
        """Get a model instance"""
        if model_name and model_name in self.models:
            return self.models[model_name]
        return self.default_model
    
    def generate_response(self, prompt: str, max_tokens: int = 4096, temperature: float = 0.7) -> str:
        """Generate response using Azure OpenAI"""
        try:
            response = self.default_model.chat.completions.create(
                model=AZURE_OPENAI_DEPLOYMENT_NAME,
                messages=[
                    {"role": "system", "content": "You are a helpful medical AI assistant. Provide accurate, evidence-based medical information while emphasizing that users should consult healthcare professionals for proper diagnosis and treatment."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return f"I apologize, but I encountered an error while processing your request. Please try again."
    
    def get_embedding(self, text: str) -> Optional[List[float]]:
        """Get embedding for text using Azure OpenAI"""
        try:
            if not self.default_model:
                logger.error("Azure OpenAI client not initialized")
                return None
                
            response = self.default_model.embeddings.create(
                model=AZURE_OPENAI_EMBEDDING_DEPLOYMENT,
                input=text
            )
            
            return response.data[0].embedding
            
        except Exception as e:
            logger.error(f"Error getting embedding: {str(e)}")
            return None


# Format markdown response with styling
def format_markdown_response(text):
    """Format the response with proper markdown and styling"""
    try:
        # Convert markdown to HTML with extras
        html = markdown2.markdown(text, extras=['fenced-code-blocks', 'tables', 'break-on-newline'])
        
        # Enhance emoji display
        emoji_map = {
            '🏥': '<span class="emoji hospital">🏥</span>',
            '💊': '<span class="emoji medication">💊</span>',
            '⚠️': '<span class="emoji warning">⚠️</span>',
            '📊': '<span class="emoji stats">📊</span>',
            '📋': '<span class="emoji clipboard">📋</span>',
            '👨‍⚕️': '<span class="emoji doctor">👨‍⚕️</span>',
            '🔬': '<span class="emoji research">🔬</span>',
            '📚': '<span class="emoji book">📚</span>',
            '🔍': '<span class="emoji search">🔍</span>',
            '🚨': '<span class="emoji alert">🚨</span>',
            '👁️': '<span class="emoji eye">👁️</span>',
            '🔄': '<span class="emoji repeat">🔄</span>',
            '🔮': '<span class="emoji crystal-ball">🔮</span>',
            '🥗': '<span class="emoji salad">🥗</span>',
            '🏋️': '<span class="emoji exercise">🏋️</span>',
            '⚖️': '<span class="emoji balance">⚖️</span>',
            '🍎': '<span class="emoji apple">🍎</span>',
            '💧': '<span class="emoji water">💧</span>'
        }
        
        for emoji, styled_emoji in emoji_map.items():
            html = html.replace(emoji, styled_emoji)
        
        return html
    except Exception as e:
        logger.error(f"Error formatting markdown: {str(e)}")
        return f"<p>Error formatting response: {str(e)}</p><pre>{text}</pre>"


# Create and initialize the agent system
agent_system = MediAssistAgentSystem()

# Global application state
APP_STATE = {
    "is_initialized": False,
    "initialization_error": None,
    "loading_progress": {
        "data": {"status": "pending", "message": "Waiting to load data"},
        "embeddings": {"status": "pending", "message": "Waiting to load embeddings"},
        "models": {"status": "pending", "message": "Waiting to initialize models"}
    },
    "start_time": None
}

# Flask routes
@app.route('/')
def home():
    """Serve the main page"""
    return render_template('index.html')

@app.route('/status', methods=['GET'])
def get_status():
    """Get application initialization status"""
    return jsonify({
        "initialized": APP_STATE["is_initialized"],
        "error": APP_STATE["initialization_error"],
        "progress": APP_STATE["loading_progress"]
    })

@app.route('/query', methods=['POST'])
def process_query():
    """Process a medical query with automatic category detection"""
    try:
        # Get query data from request
        query_data = request.json
        user_query = query_data.get('query')
        
        # Validate input
        if not user_query:
            return jsonify({
                "status": "error",
                "response": "Missing query text"
            })
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            background_thread = threading.Thread(target=agent_system.initialize)
            background_thread.start()
            
            return jsonify({
                "status": "error",
                "response": "System is initializing. Please try again in a moment."
            })
        
        # Process the query through the agent system
        # Note: No query_type is provided, so it will be automatically determined
        response = agent_system.process_query(user_query)
        
        # Format the response if successful
        if response["status"] == "success":
            html_response = format_markdown_response(response["response"])
            response["response"] = html_response
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in query endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "response": f"An error occurred: {str(e)}"
        })

@app.route('/analyze-image', methods=['POST'])
def analyze_image():
    """Analyze a medical image"""
    try:
        # Validate input
        if 'image' not in request.files or 'prompt' not in request.form:
            return jsonify({
                "status": "error",
                "response": "Missing image or prompt"
            })
        
        image_file = request.files['image']
        prompt = request.form['prompt']
        
        if image_file.filename == '':
            return jsonify({
                "status": "error",
                "response": "No image selected"
            })
        
        # Read the image data
        image_data = image_file.read()
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            background_thread = threading.Thread(target=agent_system.initialize)
            background_thread.start()
            
            return jsonify({
                "status": "error",
                "response": "System is initializing. Please try again in a moment."
            })
        
        # Process the image through the agent system
        response = agent_system.process_image_analysis(image_data, prompt)
        
        # Format the response if successful
        if response["status"] == "success":
            html_response = format_markdown_response(response["response"])
            response["response"] = html_response
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in image analysis endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "response": f"An error occurred: {str(e)}"
        })

@app.route('/analyze-reports', methods=['POST'])
def analyze_reports():
    """Analyze multiple medical reports"""
    try:
        # Validate input
        if 'reports' not in request.files:
            return jsonify({
                "status": "error",
                "response": "No reports uploaded"
            })
        
        uploaded_files = request.files.getlist('reports')
        patient_info = {}
        
        # Get patient information if provided
        if 'patient_info' in request.form:
            try:
                patient_info = json.loads(request.form['patient_info'])
            except:
                pass
        
        if not uploaded_files or all(f.filename == '' for f in uploaded_files):
            return jsonify({
                "status": "error",
                "response": "No files selected"
            })
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            background_thread = threading.Thread(target=agent_system.initialize)
            background_thread.start()
            
            return jsonify({
                "status": "error",
                "response": "System is initializing. Please try again in a moment."
            })
        
        # Process uploaded files
        reports_data = []
        for file in uploaded_files:
            if file.filename != '' and allowed_file(file.filename):
                try:
                    # Save the file temporarily
                    filename = secure_filename(file.filename)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    unique_filename = f"{timestamp}_{filename}"
                    file_path = os.path.join(app.config['REPORTS_FOLDER'], unique_filename)
                    file.save(file_path)
                    
                    # Extract content based on file type
                    file_extension = filename.rsplit('.', 1)[1].lower()
                    content = ""
                    
                    if file_extension == 'pdf':
                        content = extract_text_from_pdf(file_path)
                    elif file_extension in ['docx', 'doc']:
                        content = extract_text_from_docx(file_path)
                    elif file_extension == 'txt':
                        content = extract_text_from_txt(file_path)
                    elif file_extension in ['xlsx', 'xls']:
                        content = extract_data_from_excel(file_path)
                    elif file_extension == 'csv':
                        content = extract_data_from_csv(file_path)
                    
                    if content:
                        reports_data.append({
                            'filename': filename,
                            'content': content,
                            'size': len(content),
                            'type': file_extension
                        })
                    
                    # Clean up temporary file
                    try:
                        os.remove(file_path)
                    except:
                        pass
                        
                except Exception as e:
                    logger.error(f"Error processing file {file.filename}: {str(e)}")
                    continue
        
        if not reports_data:
            return jsonify({
                "status": "error",
                "response": "No valid reports could be processed. Please check file formats."
            })
        
        # Process the reports through the agent system
        response = agent_system.process_report_analysis(reports_data, patient_info)
        
        # Format the response if successful
        if response["status"] == "success":
            html_response = format_markdown_response(response["response"])
            response["response"] = html_response
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in report analysis endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "response": f"An error occurred: {str(e)}"
        })

@app.route('/diet-plan', methods=['POST'])
def create_diet_plan():
    """Generate a personalized diet plan"""
    try:
        # Get user data from request
        user_data = request.json
        
        # Validate required fields
        required_fields = ['age', 'gender', 'height', 'weight', 'health_goal']
        missing_fields = [field for field in required_fields if not user_data.get(field)]
        
        if missing_fields:
            return jsonify({
                "status": "error",
                "response": f"Missing required fields: {', '.join(missing_fields)}"
            })
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            background_thread = threading.Thread(target=agent_system.initialize)
            background_thread.start()
            
            return jsonify({
                "status": "error",
                "response": "System is initializing. Please try again in a moment."
            })
        
        # Process the diet plan through the agent system
        response = agent_system.process_diet_plan(user_data)
        
        # Format the response if successful
        if response["status"] == "success":
            html_response = format_markdown_response(response["response"])
            response["response"] = html_response
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in diet plan endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "response": f"An error occurred: {str(e)}"
        })

@app.route('/refresh-embeddings', methods=['POST'])
def refresh_embeddings():
    """Refresh embeddings for a specific category"""
    try:
        # Get category to refresh
        request_data = request.json or {}
        category = request_data.get('category')
        
        if not category:
            return jsonify({
                "status": "error",
                "message": "Category parameter is required"
            })
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            background_thread = threading.Thread(target=agent_system.initialize)
            background_thread.start()
            
            return jsonify({
                "status": "error",
                "message": "System is initializing. Please try again in a moment."
            })
        
        # Refresh embeddings through the agent system
        response = agent_system.refresh_embeddings(category)
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error in refresh embeddings endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"An error occurred: {str(e)}"
        })

@app.route('/templates/<path:path>')
def send_template(path):
    return send_from_directory('templates', path)

@app.route('/static/charts/<path:filename>')
def serve_chart(filename):
    """Serve generated chart files"""
    return send_from_directory(app.config['CHARTS_FOLDER'], filename)

@app.route('/get-doctors', methods=['POST'])
def get_doctors():
    """Get doctor recommendations based on symptoms"""
    try:
        # Get symptoms from request
        request_data = request.json
        symptoms = request_data.get('symptoms')
        
        if not symptoms:
            return jsonify({
                "status": "error",
                "message": "Symptoms parameter is required"
            })
        
        # Ensure system is initialized
        if not agent_system.is_initialized:
            return jsonify({
                "status": "error",
                "message": "System is initializing. Please try again in a moment."
            })
        
        # Generate specialist recommendations
        specialists_result = agent_system.generate_specialists(symptoms)
        
        return jsonify(specialists_result)
        
    except Exception as e:
        logger.error(f"Error getting doctor recommendations: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"An error occurred: {str(e)}"
        })

@app.route('/book-appointment', methods=['POST'])
def book_appointment():
    """Book an appointment with a doctor"""
    try:
        # Get booking data from request
        booking_data = request.json
        
        # Validate required fields
        required_fields = ['doctor_id', 'doctor_name', 'slot_id', 'email']
        missing_fields = [field for field in required_fields if not booking_data.get(field)]
        
        if missing_fields:
            return jsonify({
                "status": "error",
                "message": f"Missing required fields: {', '.join(missing_fields)}"
            })
        
        # Extract data
        doctor_id = booking_data['doctor_id']
        doctor_name = booking_data['doctor_name']
        slot_id = booking_data['slot_id']
        email = booking_data['email']
        
        # Get specialists data
        specialists = booking_data.get('specialists', [])
        if not specialists:
            return jsonify({
                "status": "error",
                "message": "No specialists data provided"
            })
        
        # Find the selected doctor
        selected_doctor = None
        for doctor in specialists:
            if doctor['id'] == doctor_id:
                selected_doctor = doctor
                break
                
        if not selected_doctor:
            return jsonify({
                "status": "error",
                "message": f"Doctor with ID {doctor_id} not found"
            })
            
        # Find the selected time slot
        selected_slot = None
        for slot in selected_doctor.get('available_slots', []):
            if slot_id == slot.get('formatted_time'):
                selected_slot = slot
                break
                
        if not selected_slot:
            return jsonify({
                "status": "error",
                "message": f"Time slot {slot_id} not found for doctor {doctor_id}"
            })
        
        # Check if the slot is already booked
        if selected_slot.get('is_booked', False):
            return jsonify({
                "status": "error",
                "message": "This time slot is already booked. Please select another time."
            })
        
        # Generate Microsoft Teams meeting link
        teams_meeting_id = generate_teams_meeting_id()
        teams_meeting_link = f"https://teams.microsoft.com/l/meetup-join/{teams_meeting_id}"
        
        # Mark slot as booked
        selected_slot['is_booked'] = True
        
        # Save appointment
        appointment_id = str(uuid.uuid4())
        start_time = selected_slot['start_time']
        end_time = selected_slot['end_time']

        if isinstance(start_time, str):
            try:
                # Try the original format first
                start_time = datetime.strptime(start_time, "%Y-%m-%d %H:%M")
            except ValueError:
                # If that fails, try to parse RFC format
                from email.utils import parsedate_to_datetime
                try:
                    start_time = parsedate_to_datetime(start_time)
                except:
                    # As a fallback, try another common format
                    try:
                        start_time = datetime.strptime(start_time, "%a, %d %b %Y %H:%M:%S GMT")
                    except:
                        # Last resort: create a current datetime (not ideal but prevents crashing)
                        start_time = datetime.now()

        # Similar handling for end_time
        if isinstance(end_time, str):
            try:
                end_time = datetime.strptime(end_time, "%Y-%m-%d %H:%M")
            except ValueError:
                from email.utils import parsedate_to_datetime
                try:
                    end_time = parsedate_to_datetime(end_time)
                except:
                    try:
                        end_time = datetime.strptime(end_time, "%a, %d %b %Y %H:%M:%S GMT")
                    except:
                        # Add 45 minutes to start_time as a fallback
                        end_time = start_time + timedelta(minutes=45)

        # Save appointment with datetime objects
        new_appointment = {
            "id": appointment_id,
            "doctor_id": doctor_id,
            "doctor_name": doctor_name,
            "start_time": start_time,
            "end_time": end_time,
            "formatted_time": selected_slot['formatted_time'],
            "patient_email": email,
            "teams_meeting_link": teams_meeting_link
        }
        appointments.append(new_appointment)
        
        # Send email confirmation
        email_sent = False
        try:
            email_sent = send_appointment_email(new_appointment, doctor_name, email, teams_meeting_link)
        except Exception as e:
            logger.error(f"Error sending email: {str(e)}")
        
        return jsonify({
            "status": "success",
            "message": "Appointment booked successfully",
            "appointment": {
                "id": appointment_id,
                "doctor_name": doctor_name,
                "time": selected_slot['formatted_time'],
                "teams_meeting_link": teams_meeting_link,
                "email_sent": email_sent
            }
        })
        
    except Exception as e:
        logger.error(f"Error booking appointment: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"An error occurred: {str(e)}"
        })

# Create necessary directories on startup
def create_folders():
    """Create necessary folders for the application"""
    os.makedirs('static', exist_ok=True)
    os.makedirs('templates', exist_ok=True)
    os.makedirs('data', exist_ok=True)
    os.makedirs('data/embeddings', exist_ok=True)
    os.makedirs('cache', exist_ok=True)
    os.makedirs('cache/image_analysis', exist_ok=True)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(REPORTS_FOLDER, exist_ok=True)
    os.makedirs(CHARTS_FOLDER, exist_ok=True)

# Initialize the system in background
def background_initialize():
    global APP_STATE
    
    try:
        # Create necessary folders
        create_folders()
        
        # Initialize the agent system
        success = agent_system.initialize()
        
        if success:
            APP_STATE["is_initialized"] = True
            logger.info("Background initialization completed successfully")
        else:
            APP_STATE["initialization_error"] = "Failed to initialize agent system"
            logger.error("Failed to initialize agent system")
            
    except Exception as e:
        APP_STATE["initialization_error"] = str(e)
        logger.error(f"Error in background initialization: {str(e)}")

# Start background initialization
background_thread = threading.Thread(target=background_initialize)
background_thread.daemon = True
background_thread.start()

# Run the Flask app
if __name__ == '__main__':
    # Record start time
    APP_STATE["start_time"] = time.time()
    
    # Start Flask app
    logger.info("Starting Flask application")
    app.run(debug=True, host='0.0.0.0', port=5002, threaded=True)